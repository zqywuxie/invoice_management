"""
DocxExportService for exporting invoices with expense vouchers to DOCX format.
DOCX导出服务 - 负责将发票和支出凭证导出为Word文档
"""

import io
import os
from typing import List, Optional, Tuple

from docx import Document
from docx.shared import Cm, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pdf2image import convert_from_path, convert_from_bytes
from PIL import Image

from src.models import ExpenseVoucher, Invoice
from src.sqlite_data_store import SQLiteDataStore
from src.voucher_service import VoucherService


class DocxExportService:
    """
    DOCX导出服务类，负责将发票和支出凭证导出为Word文档
    """
    
    # Standard page margins (reduced for better image display)
    PAGE_MARGIN_CM = 1.5
    
    # Voucher grid settings
    VOUCHER_COLUMNS = 2
    VOUCHER_SPACING_CM = 0.5
    
    # A4 page dimensions in cm (excluding margins)
    PAGE_WIDTH_CM = 21.0
    PAGE_HEIGHT_CM = 29.7
    
    def __init__(self, data_store: SQLiteDataStore, voucher_service: VoucherService):
        """
        初始化DOCX导出服务
        
        Args:
            data_store: SQLite数据存储实例
            voucher_service: 凭证服务实例
        """
        self.data_store = data_store
        self.voucher_service = voucher_service
    
    def convert_pdf_to_image(self, pdf_path: str = None, pdf_data: bytes = None, 
                             dpi: int = 200) -> bytes:
        """
        将PDF转换为图片
        
        Args:
            pdf_path: PDF文件路径（与pdf_data二选一）
            pdf_data: PDF二进制数据（与pdf_path二选一）
            dpi: 转换分辨率，默认200
            
        Returns:
            PNG格式的图片二进制数据
            
        Raises:
            ValueError: 未提供pdf_path或pdf_data时抛出
            RuntimeError: PDF转换失败时抛出
        """
        if pdf_path is None and pdf_data is None:
            raise ValueError("必须提供pdf_path或pdf_data参数")
        
        try:
            # 使用 use_cropbox=False 和 transparent=True 确保完整转换PDF页面
            # 不使用裁剪框，保留完整页面内容
            if pdf_data:
                images = convert_from_bytes(
                    pdf_data, 
                    dpi=dpi,
                    use_cropbox=False,
                    use_pdftocairo=True  # 使用pdftocairo可能有更好的渲染效果
                )
            else:
                images = convert_from_path(
                    pdf_path, 
                    dpi=dpi,
                    use_cropbox=False,
                    use_pdftocairo=True
                )
            
            if not images:
                raise RuntimeError("PDF转换失败：未生成任何图片")
            
            # Take the first page
            image = images[0]
            
            # Convert to bytes
            img_buffer = io.BytesIO()
            image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            return img_buffer.getvalue()
            
        except Exception as e:
            raise RuntimeError(f"PDF转换失败：{str(e)}")
    
    def _calculate_scaled_dimensions(self, original_width: int, original_height: int,
                                     max_width: float, max_height: float) -> Tuple[float, float]:
        """
        计算保持宽高比的缩放尺寸
        
        Args:
            original_width: 原始宽度（像素）
            original_height: 原始高度（像素）
            max_width: 最大宽度（厘米）
            max_height: 最大高度（厘米）
            
        Returns:
            (缩放后宽度, 缩放后高度) 单位为厘米
        """
        if original_width == 0 or original_height == 0:
            return (max_width, max_height)
        
        aspect_ratio = original_width / original_height
        
        # Try fitting by width first
        scaled_width = max_width
        scaled_height = max_width / aspect_ratio
        
        # If height exceeds max, fit by height instead
        if scaled_height > max_height:
            scaled_height = max_height
            scaled_width = max_height * aspect_ratio
        
        return (scaled_width, scaled_height)

    def _get_image_dimensions(self, image_data: bytes) -> Tuple[int, int]:
        """
        获取图片尺寸
        
        Args:
            image_data: 图片二进制数据
            
        Returns:
            (宽度, 高度) 单位为像素
        """
        img = Image.open(io.BytesIO(image_data))
        return img.size
    
    def _add_invoice_page(self, doc: Document, invoice: Invoice) -> None:
        """
        添加发票页面（第一页）
        
        Args:
            doc: Word文档对象
            invoice: 发票对象
        """
        # Get PDF data from database
        pdf_data = self.data_store.get_pdf_data(invoice.invoice_number)
        
        if pdf_data:
            # Convert PDF to image
            image_data = self.convert_pdf_to_image(pdf_data=pdf_data)
        elif invoice.file_path and os.path.exists(invoice.file_path):
            # Fallback to file path
            image_data = self.convert_pdf_to_image(pdf_path=invoice.file_path)
        else:
            # No PDF available, add placeholder text
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(f"发票号码: {invoice.invoice_number}")
            run.font.size = Pt(14)
            return
        
        # Calculate available space (page size minus margins)
        available_width = self.PAGE_WIDTH_CM - (2 * self.PAGE_MARGIN_CM)
        available_height = self.PAGE_HEIGHT_CM - (2 * self.PAGE_MARGIN_CM)
        
        # Get image dimensions and calculate scaled size
        img_width, img_height = self._get_image_dimensions(image_data)
        scaled_width, scaled_height = self._calculate_scaled_dimensions(
            img_width, img_height, available_width, available_height
        )
        
        # Add image to document
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 移除段落的左右缩进，确保图片完整显示
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(image_data), width=Cm(scaled_width))
    
    def create_voucher_grid(self, vouchers: List[ExpenseVoucher], doc: Document) -> None:
        """
        创建凭证网格布局（第二页）
        
        Args:
            vouchers: 支出凭证列表
            doc: Word文档对象
        """
        if not vouchers:
            return
        
        # Add page break
        doc.add_page_break()
        
        # Add header "支出凭证"
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("支出凭证")
        run.font.size = Pt(16)
        run.font.bold = True
        
        # Add some spacing after header
        doc.add_paragraph()
        
        # Calculate cell dimensions
        available_width = self.PAGE_WIDTH_CM - (2 * self.PAGE_MARGIN_CM)
        cell_width = (available_width - self.VOUCHER_SPACING_CM) / self.VOUCHER_COLUMNS
        
        # Calculate number of rows needed
        num_rows = (len(vouchers) + self.VOUCHER_COLUMNS - 1) // self.VOUCHER_COLUMNS
        
        # Create table for voucher grid
        table = doc.add_table(rows=num_rows, cols=self.VOUCHER_COLUMNS)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Add vouchers to table cells
        for idx, voucher in enumerate(vouchers):
            row_idx = idx // self.VOUCHER_COLUMNS
            col_idx = idx % self.VOUCHER_COLUMNS
            
            cell = table.rows[row_idx].cells[col_idx]
            
            # Load and add voucher image
            if os.path.exists(voucher.file_path):
                try:
                    with open(voucher.file_path, 'rb') as f:
                        image_data = f.read()
                    
                    # Get image dimensions
                    img_width, img_height = self._get_image_dimensions(image_data)
                    
                    # Calculate scaled dimensions to fit in cell
                    # Leave some margin within the cell
                    max_cell_width = cell_width - 0.5
                    max_cell_height = 10.0  # Max height per voucher in cm
                    
                    scaled_width, scaled_height = self._calculate_scaled_dimensions(
                        img_width, img_height, max_cell_width, max_cell_height
                    )
                    
                    # Add image to cell
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(io.BytesIO(image_data), width=Cm(scaled_width))
                    
                except Exception:
                    # If image loading fails, add placeholder
                    cell.text = f"[凭证: {voucher.original_filename}]"
            else:
                cell.text = f"[凭证文件不存在: {voucher.original_filename}]"
    
    def export_invoice_with_vouchers(self, invoice_number: str, 
                                     output_path: str = None) -> str:
        """
        导出发票和支出凭证为DOCX文档
        
        Args:
            invoice_number: 发票号码
            output_path: 输出文件路径（可选，默认生成基于发票号的文件名）
            
        Returns:
            生成的DOCX文件路径
            
        Raises:
            ValueError: 发票不存在时抛出
            RuntimeError: 文档生成失败时抛出
        """
        # Get invoice
        invoice = self.data_store.get_invoice_by_number(invoice_number)
        if not invoice:
            raise ValueError(f"发票不存在: {invoice_number}")
        
        # Create document
        doc = Document()
        
        # Set page margins (2.54cm = 1 inch)
        for section in doc.sections:
            section.top_margin = Cm(self.PAGE_MARGIN_CM)
            section.bottom_margin = Cm(self.PAGE_MARGIN_CM)
            section.left_margin = Cm(self.PAGE_MARGIN_CM)
            section.right_margin = Cm(self.PAGE_MARGIN_CM)
        
        # Add invoice on first page
        self._add_invoice_page(doc, invoice)
        
        # Get vouchers for this invoice
        vouchers = self.voucher_service.get_vouchers(invoice_number)
        
        # Add voucher grid on second page (if vouchers exist)
        if vouchers:
            self.create_voucher_grid(vouchers, doc)
        
        # Generate output path if not provided
        if output_path is None:
            # Sanitize invoice number for filename
            safe_invoice_number = invoice_number.replace('/', '_').replace('\\', '_')
            output_path = f"发票_{safe_invoice_number}.docx"
        
        # Save document
        try:
            doc.save(output_path)
        except Exception as e:
            raise RuntimeError(f"文档生成失败: {str(e)}")
        
        return output_path
    
    def export_multiple_invoices(self, invoice_numbers: List[str], 
                                  output_path: str = None) -> str:
        """
        批量导出多个发票和支出凭证为单个DOCX文档
        
        每个发票按顺序排列：第一页是发票PDF图片，第二页是凭证
        
        Args:
            invoice_numbers: 发票号码列表（按顺序）
            output_path: 输出文件路径（可选）
            
        Returns:
            生成的DOCX文件路径
            
        Raises:
            ValueError: 发票列表为空或发票不存在时抛出
            RuntimeError: 文档生成失败时抛出
        """
        if not invoice_numbers:
            raise ValueError("发票列表不能为空")
        
        # Create document
        doc = Document()
        
        # Set page margins
        for section in doc.sections:
            section.top_margin = Cm(self.PAGE_MARGIN_CM)
            section.bottom_margin = Cm(self.PAGE_MARGIN_CM)
            section.left_margin = Cm(self.PAGE_MARGIN_CM)
            section.right_margin = Cm(self.PAGE_MARGIN_CM)
        
        is_first_invoice = True
        
        for invoice_number in invoice_numbers:
            # Get invoice
            invoice = self.data_store.get_invoice_by_number(invoice_number)
            if not invoice:
                # Skip non-existent invoices but continue with others
                continue
            
            # Add page break before each invoice (except the first one)
            if not is_first_invoice:
                doc.add_page_break()
            is_first_invoice = False
            
            # Add invoice PDF image on first page
            self._add_invoice_page(doc, invoice)
            
            # Get vouchers for this invoice
            vouchers = self.voucher_service.get_vouchers(invoice_number)
            
            # Add voucher grid on second page (if vouchers exist)
            if vouchers:
                self.create_voucher_grid(vouchers, doc)
        
        # Generate output path if not provided
        if output_path is None:
            from datetime import datetime
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_path = f"发票批量导出_{timestamp}.docx"
        
        # Save document
        try:
            doc.save(output_path)
        except Exception as e:
            raise RuntimeError(f"文档生成失败: {str(e)}")
        
        return output_path
    
    def generate_export_filename(self, invoice_number: str) -> str:
        """
        生成导出文件名
        
        Args:
            invoice_number: 发票号码
            
        Returns:
            包含发票号码的文件名
        """
        # Sanitize invoice number for filename
        safe_invoice_number = invoice_number.replace('/', '_').replace('\\', '_')
        return f"发票_{safe_invoice_number}.docx"
